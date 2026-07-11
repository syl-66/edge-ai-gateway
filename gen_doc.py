#!/usr/bin/env python3
"""生成 Edge AI Agent Gateway 完整设计文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

def h(text, level=1):
    hh = doc.add_heading(text, level=level)
    for r in hh.runs: r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def para(text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text); r.font.size = Pt(11)

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

# ==================== 封面 ====================
for _ in range(6): doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Edge AI Agent Gateway"); r.bold = True; r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("边缘 AI 智能家居网关 — 完整设计文档"); r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph()
m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run("硬件平台：NXP i.MX6ULL (Cortex-A7) + Linux 4.1.15").font.size = Pt(11)
m2 = doc.add_paragraph(); m2.alignment = WD_ALIGN_PARAGRAPH.CENTER
m2.add_run("技术栈：C（嵌入式 sysfs/mmap） + Python（Agent LLM） + MQTT + SQLite3").font.size = Pt(11)

doc.add_page_break()

# ==================== 目录 ====================
h("目录", 1)
for t in [
    "1. 项目概述",
    "2. 系统架构总览",
    "3. 数据流全景 - 打开灯全链路",
    "4. Python Agent 端 - 大脑",
    "5. MQTT 通信层 - 神经",
    "6. C 嵌入式端 - 手脚",
    "  6.1 多线程架构",
    "  6.2 传感器（DHT11 / BH1750 / GP2Y1014AU / W25Q32）",
    "  6.3 执行器（LED灯 / 继电器 / 红外 / 状态LED）",
    "  6.4 GPIO 访问策略（sysfs vs mmap）",
    "  6.5 工具分发",
    "  6.6 本地存储 SQLite3",
    "7. 通信协议 JSON-RPC",
    "8. 目录结构",
    "9. 面试要点",
]: para(t)

doc.add_page_break()

# ==================== 第1章 ====================
h("1. 项目概述", 1)
para('这是一个面向嵌入式 Linux 应用开发岗位的综合性项目。用户用自然语言发出指令（如"打开灯"），AI Agent 理解意图后通过 MQTT 下发指令到 i.MX6ULL 开发板，C 程序操作 GPIO 控制真实的继电器、LED 和传感器。')
doc.add_paragraph()
para("一句话：Python 做 AI Agent（大脑），MQTT 做通信（神经），C 做嵌入式控制（手脚）。")
doc.add_paragraph()
para("解决的问题：")
bullet("展示多语言：Python（Agent）+ C（嵌入式）+ JSON（协议）")
bullet("展示嵌入式 Linux 全栈：GPIO sysfs/mmaps、I2C、ADC、SPI、多线程、SQLite3")
bullet("展示架构设计：三层架构、消息队列、工具分发模式")
bullet("展示工程选型：GP2Y1014AU 替代 PMS5003 降成本、sysfs 标准 GPIO")
bullet("展示 AI 集成：LLM Function Calling、自动巡检、规则引擎")

doc.add_page_break()

# ==================== 第2章 架构 ====================
h("2. 系统架构总览", 1)
para("三层架构：")

arch = """+==================================================================+
|                    PC / 云服务器                                 |
|  [Web Dashboard]  [Agent Core]  [自动化规则引擎]                 |
|   Flask+SocketIO   LLM对话+       定时检查传感器                  |
|                    Function       触发联动动作                    |
|  用户打字/点击      Calling                                        |
|       |                |                |                         |
|       +----------------+----------------+                         |
|                        |  MQTT Bridge (paho-mqtt)                 |
+========================|==========================================+
                         |  MQTT (mosquitto broker)
     edge/.../tool/call     (下发指令)
     edge/.../sensor/report (传感器上报)
     edge/.../tool/result   (执行结果)
                         |
+========================|==========================================+
|                   i.MX6ULL 嵌入式板                               |
|                        |                                          |
|  +---------------------+---------------------------------------+  |
|  |  main.c (4线程)                                              |  |
|  |  [sensor_thread 2s]  [mqtt_rx]  [mqtt_tx]  [heartbeat 30s] |  |
|  |       |                  |           |            |          |  |
|  |  DHT11/BH1750/      解析JSON    消费内部队列   心跳上报       |  |
|  |  GP2Y1014AU/W25Q32   查表分发    mqtt_publish               |  |
|  |       |                  |           |            |          |  |
|  |       +--------+---------+-----------+------------+          |  |
|  |                |  内部消息队列 (环形缓冲区)                    |  |
|  |                |  pthread_mutex + pthread_cond                |  |
|  |                |                                              |  |
|  |  +-------------+-------------+                                |  |
|  |  | GPIO 访问层 (分层设计)     |                                |  |
|  |  |  /dev/mem+mmap   sysfs   |                                |  |
|  |  |  DHT11           继电器   |                                |  |
|  |  |  IR发射           LED灯   |                                |  |
|  |  |                  状态LED  |                                |  |
|  |  |                  GP2Y LED |                                |  |
|  |  +---------------------------+                                |  |
|  |                                                                |  |
|  |  SQLite3(WAL)  I2C(BH1750)  ADC(GP2Y)  SPI(W25Q32)           |  |
|  +----------------------------------------------------------------|  |
+====================================================================+"""

for line in arch.split('\n'):
    p = doc.add_paragraph(); r = p.add_run(line)
    r.font.name = 'Consolas'; r.font.size = Pt(7)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

doc.add_page_break()

# ==================== 第3章 数据流 ====================
h('3. 数据流全景 - 一条"打开灯"指令的完整旅程', 1)
para("这是整个项目最重要的章节。理解这条链路，面试就能从头讲到尾。")

steps = [
    ("第1步：用户输入",
     '用户在浏览器输入"打开灯"，点击发送\n'
     "→ web/dashboard.py → on_chat() / POST /api/chat\n"
     '→ 数据：{"message":"打开灯"}'),

    ("第2步：LLM 理解意图",
     "EdgeAgent.chat() → 调用 DeepSeek API\n"
     "→ 传入 TOOLS_OPENAI_FORMAT 工具列表\n"
     '→ LLM 返回 tool_calls: {name:"control_led", arguments:{action:"on"}}\n'
     "→ agent_core.py → chat()"),

    ("第3步：工具执行",
     "execute_tool() 查 TOOL_REGISTRY\n"
     "→ control_led(args) → mqtt_bridge.call_tool()\n"
     "→ tools/__init__.py → control_led()"),

    ("第4步：MQTT 下发",
     'MqttBridge 发布: topic=edge/imx6ull-gateway-001/tool/call\n'
     'payload={"id":"req_...","method":"tool.call","params":{"tool":"control_led","args":{"action":"on"}}}\n'
     "→ mqtt_bridge.py → call_tool()"),

    ("第5步：C端接收",
     "libmosquitto on_message → mqtt_msg_enqueue（环形缓冲区）\n"
     "→ mqtt_rx_thread 取出 → tool_call_parse_json() 解析\n"
     "→ mqtt_client.c + tool_dispatcher.c"),

    ("第6步：工具分发+硬件控制",
     'tool_dispatch() 查 g_tool_registry[] → "control_led" 匹配\n'
     "→ tool_control_led(args, result, len)\n"
     '→ led_control_set("on") → gpio_write_val(4, 1)\n'
     "→ write /sys/class/gpio/gpio4/value → GPIO1_IO04 输出 3.3V\n"
     "→ 220Ω 电阻 → LED 导通 → 灯亮\n"
     "→ tool_handlers.c → led_control.c → gpio_util.c"),

    ("第7步：结果回传",
     "tool_result_to_json() 序列化 → iq_put(&g_tool_iq)\n"
     "→ mqtt_tx_thread 消费 → mqtt_publish()\n"
     '→ topic: edge/.../tool/result, payload: {"success":true}\n'
     "→ main.c → mqtt_tx_thread"),

    ("第8步：Agent 收尾",
     "on_tool_result() → Event.set() → 工具结果加入 LLM 上下文\n"
     '→ LLM 再次推理 → "灯已打开"\n'
     "→ WebSocket 推送到浏览器\n"
     "→ tools/__init__.py → on_tool_result()"),
]
for title, cont in steps:
    para(title)
    for run in doc.paragraphs:
        if title in (run.text or ''): run.runs[0].bold = True
    para(cont)

doc.add_page_break()

# ==================== 第4章 ====================
h("4. Python Agent 端 - 大脑", 1)

h("4.1 核心模块", 2)
bullet("agent_core.py — EdgeAgent：LLM 对话 + Function Calling 循环 + 传感器缓存 + 主动巡检")
bullet("mqtt_bridge.py — MqttBridge：MQTT 连接/发布/订阅/事件路由/工具调用下发")
bullet("tools/__init__.py — 8 个工具函数：control_led/control_relay/send_ir/read_temperature/read_humidity/get_device_status/get_weather/create_automation")
bullet("web/dashboard.py — Flask + SocketIO Web 面板")
bullet("memory/context_store.py — 对话记忆 + 传感器趋势")

h("4.2 Function Calling 机制", 2)
para("这是整个 AI Agent 的核心设计模式：LLM 不直接控制硬件，而是输出函数调用，由 Python 实际执行。")
code('用户: "把灯打开"')
code('  → LLM(DeepSeek): tool_calls[{name:"control_led", arguments:{action:"on"}}]')
code("  → execute_tool('control_led', {action:'on'})")
code("  → MQTT 下发 → 硬件执行 → LLM 生成回复")
para("关键代码在 agent_core.py 的 chat()：最多 5 轮工具调用循环，防止 LLM 死循环。")

h("4.3 为什么不用 LangChain", 2)
para("LangChain 封装太深。直接用 OpenAI SDK 写 Function Calling 循环，50 行代码说清楚，面试能逐行解释。")

doc.add_page_break()

# ==================== 第5章 ====================
h("5. MQTT 通信层 - 神经", 1)

h("5.1 为什么选 MQTT", 2)
bullet("长连接：嵌入式资源受限，HTTP 反复建连开销大")
bullet("发布/订阅：Agent 和设备解耦，QoS=1 消息不丢")
bullet("双向：同一条连接上报传感器 + 下发指令")
bullet("低带宽：最小帧头 2 字节")

h("5.2 Topic 设计", 2)
code("设备上报传感器:  edge/imx6ull-gateway-001/sensor/report")
code("Agent下发工具调用: edge/imx6ull-gateway-001/tool/call")
code("设备回传结果:     edge/imx6ull-gateway-001/tool/result")
code("设备心跳:         edge/imx6ull-gateway-001/heartbeat")
code("设备注册(retained):edge/imx6ull-gateway-001/register")

h("5.3 双端实现对比", 2)
code("Python端 (mqtt_bridge.py):   C端 (mqtt_client.c):")
code("  paho-mqtt                    libmosquitto")
code("  on_message → method路由       on_message → 环形缓冲区 → 轮询")

doc.add_page_break()

# ==================== 第6章 ====================
h("6. C 嵌入式端 - 手脚", 1)
para("嵌入式端是项目技术深度的核心 —— 纯 C 语言在 i.MX6ULL 上实现多线程、多协议传感器、GPIO 控制和本地存储。")

h("6.1 多线程架构", 2)
code("main()")
code("  ├── sensor_thread()    传感器采集  2s周期")
code("  │   DHT11→BH1750→GP2Y1014AU→SQLite→JSON→内部队列")
code("  ├── mqtt_rx_thread()   MQTT接收    阻塞轮询")
code("  │   环形缓冲区→解析JSON→tool_dispatch()→内部队列")
code("  ├── mqtt_tx_thread()   MQTT发送    消费内部队列")
code("  └── heartbeat_thread() 心跳上报   30s周期")

para("线程间通信：自研轻量消息队列 (iqueue_t)")
bullet("环形缓冲区 char[64][2048] + head/tail 指针")
bullet("pthread_mutex + pthread_cond（生产者-消费者模型）")
bullet("为什么不用 POSIX mqueue：很多嵌入式内核没开 CONFIG_MQUEUE")

h("6.2 传感器模块", 2)
para("支持 4 种传感器，覆盖 4 种接口协议：")

para("① DHT11 温湿度 — GPIO 单总线 bit-banging")
bullet("/dev/mem + mmap 直接操作 GPIO1 寄存器 (0x0209C000)")
bullet("理由：需分辨 26us vs 70us 脉宽，sysfs write 延迟 ~50us 无法做 bit-banging")
bullet("时序：clock_gettime(CLOCK_MONOTONIC) 纳秒时间戳 + busy-wait + __asm__(\"nop\")")
bullet("流程：主机拉低 20ms → 拉高 30us → 切输入 → 读 40bit → 校验 → 解析")

para("② BH1750 光照 — I2C")
bullet("open(\"/dev/i2c-0\") + ioctl(I2C_SLAVE, 0x23) + write/read")
bullet("发 0x10 → 等 180ms → 读 2 字节 → lux = raw / 1.2")

para("③ GP2Y1014AU 粉尘 — GPIO sysfs + ADC")
bullet("夏普红外粉尘传感器，成本约 2 元（PMS5003 要 30-50 元，选它降成本）")
bullet("GPIO sysfs 控制红外 LED 脉冲：gpio_export_out(9) → gpio_write_val")
bullet("ADC 读模拟电压：/sys/bus/iio/devices/iio:device0/in_voltage1_raw")
bullet("时序：LED ON → 等 0.28ms → 读 ADC → LED OFF → 等 9.68ms → 周期完成")
bullet("算法：粉尘(ug/m3) = (V_led_on - V_led_off) * 172")
bullet("接线：GPIO1_IO09 → NPN基极(1kΩ) → 集电极接150Ω到5V → 发射极接LED-GND")

para("④ W25Q32 NOR Flash — SPI")
bullet("open(\"/dev/spidev0.0\") + ioctl 配置 mode/speed/bits")
bullet("全双工：ioctl(SPI_IOC_MESSAGE(1)) 发N字节同时收N字节")
bullet("读 JEDEC ID：发 0x9F + 3字节 dummy → 收 制造商/类型/容量")

para("统一管理：sensor_manager.c")
bullet("sensor_manager_init()：逐个初始化，失败不致命")
bullet("sensor_read_all()：读所有传感器 → sensor_value_t[]")
bullet("全局缓存：g_temperature/g_humidity/g_lux/g_pm25")
bullet("sensor_report_to_json() → MQTT 上报")

h("6.3 执行器模块", 2)

para("① LED 灯 (led_control.c)")
bullet("GPIO1_IO04 → 220Ω 电阻 → LED → GND，3.3V 高电平点亮")
bullet("使用 GPIO sysfs：gpio_export_out(4) + gpio_write_val(4, 0/1)")
bullet("支持 on/off/toggle 三种操作")
bullet("工具接口：Agent 调用 control_led → MQTT → tool_control_led()")

para("② 继电器 (relay_control.c)")
bullet("GPIO sysfs → 三极管驱动 → 继电器线圈 → 触点控制风扇")
bullet("RELAY_FAN_GPIO=5 → 风扇")

para("③ 红外发射 (ir_control.c)")
bullet("NEC 协议：9ms 引导 + 32bit 数据 LSB first")
bullet("38kHz 载波：GPIO sysfs gpio_write_val() + usleep 生成")
bullet("注意：usleep 精度不足以完美产生 38kHz，实际载波频率会有抖动，适合红外控制（精度要求不高）")

para("④ 状态 LED (status_led.c) — 占位")
bullet("TODO：常亮=正常 / 慢闪=传感器异常 / 快闪=网络断开")

para("统一管理：device_manager.c")
bullet("device_manager_init() → relay + led + ir 逐个初始化")
bullet("device_manager_cleanup() → 逐个释放")

h("6.4 GPIO 访问策略", 2)
para("这是面试中体现工程决策能力的关键：")

code("+-----------------+--------------+-----------------------------+")
code("| 外设            | 访问方式     | 原因                         |")
code("+-----------------+--------------+-----------------------------+")
code("| DHT11 (温湿度)  | /dev/mem mmap| 单总线, 需分辨 26us vs 70us  |")
code("| IR 发射         | GPIO sysfs   | 38kHz载波, 红外精度要求不高    |")
code("| 继电器/LED/粉尘 | GPIO sysfs   | 低频开关, sysfs 直接写 value   |")
code("| 状态 LED        | GPIO sysfs   | 标准 GPIO 输出                |")
code("+-----------------+--------------+-----------------------------+")

para("为什么用 sysfs（而非 libgpiod）：")
bullet("sysfs 简单直接：write /sys/class/gpio/export → write direction → write value")
bullet("不需要交叉编译额外库（libgpiod 需要 libgpiod.so），减少依赖")
bullet("内核 4.x 广泛支持，兼容性好")
bullet("缺点：两步操作非原子、进程崩溃可能残留 export 状态")

para("为什么 DHT11 不用 sysfs 而用 mmap：")
bullet("DHT11 协议需要分辨 26us vs 70us 脉宽，sysfs write 延迟 ~50us 做不到")
bullet("/dev/mem + mmap 直接寄存器操作延迟 < 0.1us")
bullet("项目采用分级策略：能用 sysfs 的用 sysfs（简单通用），必须高速的用 mmap")

h("6.5 工具分发", 2)
para("嵌入式端工具分发和 Python 端 Function Calling 对称设计：")

code("Python: TOOL_REGISTRY['control_led']  C: g_tool_registry[] = {")
code("  → control_led(args)                   'control_led', tool_control_led,")
code("                                         'control_relay', tool_control_relay,")
code("    查表 → 函数调用                       'send_ir', tool_send_ir, ...}")
code("                                                 查表 → 函数指针调用")

bullet("tool_call_parse_json()：手写字符串解析，不依赖 cJSON")
bullet("tool_dispatch()：线性查表 g_tool_registry[] → 匹配函数指针 → 执行")
bullet("tool_result_to_json()：序列化为 JSON-RPC 格式")

h("6.6 本地存储 SQLite3", 2)
bullet("选 SQLite3 而非纯文件：支持 SQL 查询，不用自己写索引")
bullet("WAL 模式：读写不互斥，采集线程写时查询线程也能读")
bullet("Prepared Statement：? 占位符防注入 + 预编译提升性能")
bullet("自动清理：DELETE 超过 7 天的数据，防止磁盘写满")
bullet("SQLite3 静态编译进程序 (lib/sqlite3.c)，不依赖外部 .so")

doc.add_page_break()

# ==================== 第7章 ====================
h("7. 通信协议 JSON-RPC", 1)
para("① 工具调用 (Agent → 设备)")
code('{"id":"req_001","method":"tool.call","params":{"tool":"control_led","args":{"action":"on"}}}')
para("② 工具结果 (设备 → Agent)")
code('{"id":"req_001","method":"tool.result","result":{"success":true,"data":{...}}}')
para("③ 传感器上报 (设备 → Agent, 周期2s)")
code('{"id":"evt_001","method":"sensor.report","params":{"timestamp":"...","sensors":[{...}]}}')
para("④ 心跳 (设备 → Agent, 30s)")
code('{"id":"hb_001","method":"heartbeat","params":{"uptime":3600}}')
para("⑤ 设备注册 (retained消息)")
code('{"id":"reg_001","method":"device.register","params":{"device_id":"...","tools":["control_led",...]}}')

doc.add_page_break()

# ==================== 第8章 目录 ====================
h("8. 项目目录结构", 1)
code("edge-ai-agent-gateway/")
code("├── agent/                     Python AI Agent")
code("│   ├── main.py               启动入口")
code("│   ├── agent_core.py         EdgeAgent LLM对话+Function Calling")
code("│   ├── mqtt_bridge.py        MQTT连接+发布订阅+事件路由")
code("│   ├── config.yaml           配置")
code("│   ├── tools/                工具实现 (Agent的手)")
code("│   ├── memory/               对话记忆")
code("│   └── web/                  Flask+SocketIO面板")
code("├── embedded/                 C 嵌入式程序 (i.MX6ULL)")
code("│   ├── Makefile              交叉编译")
code("│   ├── CMakeLists.txt        CMake备选")
code("│   ├── include/              头文件")
code("│   │   ├── config.h          全局配置 (引脚/设备路径)")
code("│   │   ├── logging.h         日志宏定义")
code("│   │   ├── protocol.h        JSON-RPC协议结构体")
code("│   │   ├── actuator/         执行器头文件")
code("│   │   │   ├── gpio_util.h   GPIO sysfs操作")
code("│   │   │   ├── relay_control.h 继电器控制")
code("│   │   │   ├── led_control.h LED灯控制")
code("│   │   │   ├── ir_control.h  NEC红外发射")
code("│   │   │   └── device_manager.h 执行器统一管理")
code("│   │   ├── sensor/           传感器头文件")
code("│   │   │   ├── dht11.h       DHT11温湿度")
code("│   │   │   ├── bh1750.h      BH1750光照")
code("│   │   │   ├── gp2y1014au.h  GP2Y1014AU粉尘")
code("│   │   │   ├── w25q32.h      W25Q32 SPI Flash")
code("│   │   │   ├── status_led.h  状态LED")
code("│   │   │   └── sensor_manager.h 传感器统一管理")
code("│   │   ├── comm/             通信头文件")
code("│   │   ├── tools/            工具分发头文件")
code("│   │   └── storage/          SQLite3头文件")
code("│   └── src/                  源文件 (与include/一一对应)")
code("├── config/                   部署配置 (systemd/mosquitto)")
code("└── tests/                    Python端测试")

doc.add_page_break()

# ==================== 第9章 ====================
h("9. 面试要点", 1)

para("① 一句话定调（10秒）")
para('"这是一个边缘 AI 智能家居网关。用户说打开灯，Python Agent 通过 LLM 理解意图，MQTT 下发指令到 i.MX6ULL，C 程序操作 GPIO 控制真实的硬件。"')

para("② 展示架构（30秒）")
para('"三层：Python AI 大脑 + MQTT 通信 + C 嵌入式手脚。设备端 4 线程 + 4 种传感器协议 + GPIO sysfs + mmap。"')

para("③ 按方向展开（2分钟）")
para("偏软件：讲 Function Calling 机制 → 多线程 → 内部队列 → JSON-RPC 协议")
para("偏硬件：讲 GPIO 分级策略（mmap 用于 DHT11，sysfs 用于常规 IO）→ GP2Y1014AU 选型降成本 → 4 种外设协议")
para("偏系统：讲 SQLite3 WAL → 自研消息队列 → 为什么不用 POSIX mqueue")

para("④ 主动说坑（加分）")
bullet("GP2Y1014AU 替代 PMS5003：从 50 元激光传感器换成 2 元红外传感器降成本，但精度从 ±10% 降到 ±20%，AI 补偿：LLM 结合室外气象数据交叉校准")
bullet("IR 发射 usleep 精度不够生成 38kHz：红外遥控对载波频率精度要求不高（±2kHz 都能解码），sysfs 方案实测可用")
bullet("DHT11 mmap 资源泄漏：初始版本忘记 munmap，后来加了 cleanup 函数")

para("⑤ 如果重来")
bullet("DHT11 应改用内核驱动而非 userspace mmap")
bullet("IR 应改用硬件 PWM 生成 38kHz 载波")
bullet("GP2Y1014AU 的 ADC 应改用 IIO buffer 模式而非每次 sysfs 读")
bullet("JSON 解析应引入 cJSON 避免手写边界 bug")

para("—— 全文完 ——")

# 保存
output_path = 'Edge AI Agent Gateway 设计文档 v3.docx'
doc.save(output_path)
print(f"文档已生成: {output_path}")
